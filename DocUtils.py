import os
import time

from docx import Document
from win32com import client as wc


def doc_2_docx(old_doc, new_doc):
    word = wc.Dispatch('Word.Application')
    for parent, directory, files in os.walk(old_doc):
        for f in files:
            if not f.endswith(".doc"):
                continue
            doc = word.Documents.Open(os.path.join(parent, f))  # 目标路径下的文件
            new_file_path = os.path.join(new_doc, f.split(".")[0] + ".docx")
            doc.SaveAs(new_file_path, 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
            doc.Close()
    word.Quit()


def get_table_text_list(file_path):
    doc_str = Document(file_path)
    num_tables = doc_str.tables
    table = num_tables[0]
    row_count = len(table.rows)
    col_count = len(table.columns)
    text_list = []
    for i in range(row_count):
        for j in range(col_count):
            text_list.append(table.cell(i, j).text)
    text_list = [i.strip() for i in text_list]
    return text_list


def get_textbox_text_list(file_path):
    doc = Document(file_path)
    children = doc.element.body.iter()
    child_iters = []
    for child in children:
        # 通过类型判断目录
        if child.tag.endswith('textbox'):
            for ci in child.iter():
                if ci.tag.endswith('main}r'):
                    child_iters.append(ci)
    textbox = [ci.text for ci in child_iters]
    return textbox


def test():
    doc_2_docx(os.getcwd() + "doc", os.getcwd() + "docx")
    table_list = get_table_text_list("docx/EMKNGBKHI20460063 格式件1.docx")
    textbox_list = get_textbox_text_list("docx/EMKNGBKHI20460063 格式件1.docx")
    for index, doc in enumerate(table_list):
        print('{} : {}'.format(index, doc))
    for index, doc in enumerate(textbox_list):
        print('{} : {}'.format(index, doc))


if __name__ == '__main__':
    test()

